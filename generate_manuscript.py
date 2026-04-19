"""
Generate the full manuscript as a .docx file in npj Precision Oncology Review format.

npj PO Review requirements:
- Title: up to 15 words
- Abstract: up to 70 words, no subheadings
- Main text: 3,000-4,000 words with subheadings
- References: limited to 60
- Figure legends: up to 350 words per figure
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(0)

h1 = doc.styles['Heading 1']
h1.font.name = 'Arial'
h1.font.size = Pt(14)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0, 0, 0)
h1.paragraph_format.space_before = Pt(24)
h1.paragraph_format.space_after = Pt(12)

h2 = doc.styles['Heading 2']
h2.font.name = 'Arial'
h2.font.size = Pt(12)
h2.font.bold = True
h2.font.italic = True
h2.font.color.rgb = RGBColor(0, 0, 0)
h2.paragraph_format.space_before = Pt(18)
h2.paragraph_format.space_after = Pt(8)


def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.line_spacing = 2.0
    p.add_run(text)
    return p


def ref_entry(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.size = Pt(9)
    return p


# ============================================================
# TITLE PAGE
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
run = p.add_run(
    'Multi-Omics and AI in Lung Cancer:\n'
    'Gaps in Integration, Stratification, and Translation'
)
run.font.size = Pt(16)
run.bold = True

# Authors
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
run = p.add_run('Vignesh Ravichandran')
run.font.size = Pt(12)
run.bold = True
sup = p.add_run('1')
sup.font.superscript = True
sup.font.size = Pt(12)
p.add_run(', Suchismita Ray').font.size = Pt(12)
sup2 = p.add_run('1')
sup2.font.superscript = True
sup2.font.size = Pt(12)

# Affiliation
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
sup = p.add_run('1')
sup.font.superscript = True
sup.font.size = Pt(10)
run = p.add_run('Digital Health Institute, Rutgers University, New Brunswick, NJ, USA')
run.font.size = Pt(10)
run.italic = True

# Correspondence
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
run = p.add_run('Correspondence: vignesh.ravichandran@rutgers.edu')
run.font.size = Pt(10)

doc.add_page_break()

# ============================================================
# ABSTRACT  (max 70 words, no subheadings)
# ============================================================
doc.add_heading('Abstract', level=1)

body(
    'Lung cancer remains the leading cause of cancer death, yet multi-omics integration '
    'and AI/ML approaches have not achieved clinical translation. We review 66,879 papers '
    'across 12 themes and census 25 integration studies, identifying critical gaps: survival-based '
    'feature selection that collapses biological heterogeneity, absent sex- and never-smoker-stratified '
    'analyses, TCGA data monoculture, and missing modalities. We propose a hierarchical '
    'framework bridging molecular taxonomy with clinical utility.'
)
# 63 words

doc.add_page_break()

# ============================================================
# MAIN TEXT (~3,500 words target)
# ============================================================

# --- INTRODUCTION ---
doc.add_heading('Introduction', level=1)

body(
    'Lung cancer accounts for approximately 1.8 million deaths annually worldwide. Non-small cell '
    'lung cancer (NSCLC) comprises ~85% of cases, with lung adenocarcinoma (LUAD) and squamous '
    'cell carcinoma (LUSC) as predominant subtypes. Despite advances in targeted therapies and '
    'immune checkpoint blockade, 5-year survival remains below 25%, reflecting the inadequacy of '
    'current molecular classifications to capture the biological complexity driving treatment '
    'response and clinical outcomes.'
)

body(
    'The Cancer Genome Atlas (TCGA) established foundational molecular landscapes for LUAD [1] '
    'and LUSC [2], revealing them as distinct diseases. The Clinical Proteomic Tumor Analysis '
    'Consortium (CPTAC) demonstrated that protein-level alterations frequently diverge from '
    'genomic predictions, identifying druggable targets invisible to DNA sequencing [3-5]. '
    'Small cell lung cancer (SCLC) has been reclassified into four transcription factor-defined '
    'subtypes with distinct vulnerabilities [6,7]. These advances have expanded actionable targets '
    'but revealed the inadequacy of single-omics approaches.'
)

body(
    'Multi-omics integration -- simultaneous analysis of genomic, transcriptomic, epigenomic, and '
    'proteomic data -- promises to overcome these limitations [8,9]. Concurrently, AI/ML has been '
    'applied across the lung cancer continuum, from histopathology-based molecular prediction [10] '
    'to multi-modal survival modeling [11]. Yet translation into clinical practice has been '
    'hampered by reproducibility concerns, algorithmic bias, and absent prospective validation. '
    'Here, we review the multi-omics and AI/ML landscape in lung cancer, drawing on a systematic '
    'search of 66,879 papers across 12 themes and a focused census of 25 integration studies, '
    'to identify critical methodological gaps and propose solutions.'
)

# --- MOLECULAR HETEROGENEITY ---
doc.add_heading('Molecular heterogeneity of lung cancer', level=1)

body(
    'TCGA integrated mRNA, miRNA, methylation, and protein data from 230 LUAD tumors, establishing '
    'EGFR, KRAS, and ALK as principal drivers and defining three transcriptional subtypes [1]. '
    'The LUSC characterization revealed TP53 mutations in >90% of cases, FGFR1/SOX2 amplifications, '
    'and four molecular subtypes [2]. Co-mutation patterns -- particularly KRAS/STK11 and '
    'KRAS/KEAP1 -- define biologically distinct subgroups with differential immunotherapy '
    'sensitivity [12]. CPTAC proteogenomic studies identified therapeutic vulnerabilities through '
    'phosphoproteomic profiling that genomics alone would have missed [3-5]. East Asian populations '
    'demonstrated unique EGFR-pathway rewiring in never-smoker LUAD [5].'
)

body(
    'SCLC has been reclassified by dominant transcription factors: ASCL1, NEUROD1, POU2F3, and '
    'YAP1 [6], validated by immunohistochemistry [7] and recent proteogenomic characterization [13]. '
    'Histologic transformation, particularly NSCLC-to-SCLC conversion under TKI pressure, has '
    'emerged as a mechanism of acquired resistance [14]. Rare subtypes remain molecularly '
    'understudied [15]. Critically, most large-scale studies were conducted in European-ancestry '
    'cohorts, and sex-stratified characterization has rarely been performed.'
)

# --- MULTI-OMICS METHODS + FEATURE SELECTION ---
doc.add_heading('Multi-omics integration and the feature selection problem', level=1)

doc.add_heading('Integration frameworks', level=2)

body(
    'Integration methods have evolved from iCluster joint latent variable modeling [8] through '
    'Similarity Network Fusion (SNF) [9], MOFA+ Bayesian factor analysis [16], to MOVICS '
    'consensus clustering across ten algorithms [17]. Benchmarking reveals no single method '
    'dominates [18,19]. Deep learning approaches including variational autoencoders and graph '
    'neural networks show promise but require larger datasets [20,21]. Methods for spatial or '
    'temporal multi-omics integration remain nascent [22].'
)

doc.add_heading('The feature selection philosophy divide', level=2)

body(
    'Our systematic census of 25 multi-omics lung cancer studies reveals that feature selection '
    'strategy -- not integration algorithm -- is the most consequential methodological choice. '
    'Of 18 MOVICS-based studies (72%), 13 employ supervised Cox regression filtering, retaining '
    '~2,000-3,500 survival-associated features and invariably yielding 2 subtypes (high-risk vs. '
    'low-risk). Only 3 studies use unsupervised variance-based selection with ~10,000 features, '
    'yielding 4-5 biologically distinct subtypes with unique molecular hallmarks and therapeutic '
    'vulnerabilities [23].'
)

body(
    'This divide has profound consequences. Prognosis-first approaches collapse biologically '
    'distinct subtypes sharing similar survival into single categories, losing mechanistic '
    'information needed for precision therapy. Immune-excluded and metabolically dysfunctional '
    'tumors may both have poor prognosis but require fundamentally different treatments. Zou et '
    'al. demonstrated that variance-based clustering reveals 4 LUAD subtypes with distinct drug '
    'sensitivities invisible under binary classification [23]. Furthermore, clinical features '
    '(age, stage, smoking status, sex) are almost never integrated as input features, and no '
    'study tests whether molecular features add incremental value over clinical predictors alone. '
    'Feature selection is predominantly univariate, ignoring multivariate approaches such as '
    'elastic net, WGCNA modules, or stability selection [24]. Additionally, 88% of studies rely '
    'on TCGA, zero use DIABLO despite its success in breast cancer, and zero include metabolomics '
    'or spatial omics [24].'
)

# --- AI/ML ---
doc.add_heading('AI/ML in lung cancer', level=1)

body(
    'Deep learning can classify NSCLC subtypes from H&E-stained whole slide images with '
    'pathologist-level accuracy and predict genomic alterations directly from tissue morphology '
    '[10]. Multi-modal frameworks integrating histology with genomics outperform single-modality '
    'approaches [11]. Foundation models pre-trained on massive pathology datasets have emerged as '
    'transformative since 2023, enabling fine-tuning for diverse tasks with minimal labeled '
    'data [25,26]. Radiomics-based immunotherapy prediction models show promise but face '
    'reproducibility challenges across scanners and institutions [27].'
)

body(
    'External validation studies consistently demonstrate performance degradation, and algorithmic '
    'bias from demographic underrepresentation is a systemic concern [28]. Few models have been '
    'validated on multi-ethnic cohorts, sex-stratified evaluation is virtually absent, and AI '
    'tools for never-smoker NSCLC do not exist. The convergence of AI with multi-omics data '
    'represents the next frontier, yet this integration remains methodologically unstandardized.'
)

# --- SEX, SMOKING, ANCESTRY ---
doc.add_heading('Sex, smoking, and ancestry stratification', level=1)

body(
    'Women show higher EGFR mutation rates in LUAD, particularly among never-smokers, and estrogen '
    'receptor signaling modulates lung cancer biology [29,30]. Meta-analyses suggest differential '
    'immunotherapy benefit by sex, though findings are inconsistent [31,32]. Sex-stratified drug '
    'response data are absent from major pharmacogenomic databases. Never-smoker lung cancer '
    'comprises three molecular subtypes -- piano (UBA1-driven), mezzo (EGFR-dominant), and forte '
    '(KRAS-dominant) -- with distinct evolutionary trajectories [33]. The tumor microenvironment '
    'tends to be immunologically "cold" [34], complicating immunotherapy selection.'
)

body(
    'PM2.5 promotes EGFR-driven LUAD through inflammatory IL-1-beta signaling that expands '
    'pre-existing mutant clones rather than causing new mutations -- a paradigm shift from '
    'mutagenic to promotional carcinogenesis [14]. Radon carries dose-dependent risk with '
    'distinct molecular signatures [35,36]. The exposome-to-subtype intersection remains '
    'virtually unmapped: we know environmental exposures cause lung cancer but not which molecular '
    'subtypes they produce.'
)

# --- EPIGENETICS ---
doc.add_heading('Epigenetic alterations', level=1)

body(
    'DNA methylation changes are early carcinogenic events, with hypermethylation of CDKN2A, '
    'RASSF1A, and MGMT reproducible across studies [37]. Chromatin remodeling plays mechanistic '
    'roles: KMT2D deficiency creates glycolytic vulnerabilities [38], SWI/SNF complexes promote '
    'TKI resistance [39], and NSD3 methylation drives squamous cell lung cancer [40]. cfDNA '
    'methylation enables SCLC subtype identification [41] and brain metastasis prediction [42]. '
    'Notably, 76 retracted papers in our epigenetics collection -- concentrated in miRNA/lncRNA '
    'prognostic studies -- signal reproducibility concerns requiring pre-registered validation.'
)

# --- IMMUNE BIOMARKERS ---
doc.add_heading('Immune biomarkers and immunotherapy', level=1)

body(
    'TMB associates with PD-1 blockade sensitivity [43] but molecular context modifies utility: '
    'STK11/KEAP1 co-mutations confer resistance despite high TMB [12]. Pan-cancer immune '
    'landscape analyses defined immune subtypes [44], computational immunogenicity scores [45], '
    'and conserved TME archetypes predicting immunotherapy response [46]. The gut microbiome '
    'modulates immunotherapy efficacy [47,48]. Despite this rich literature, composite biomarker '
    'panels remain unstandardized and spatial immune biomarkers lack prospective validation.'
)

# --- TRANSLATIONAL + DRUG REPURPOSING ---
doc.add_heading('Translational bottlenecks and drug repurposing', level=1)

body(
    'Biomarker testing rates remain suboptimal with disparities by race, socioeconomic status, '
    'and geography [49,50]. Molecular tumor boards improve treatment selection [51,52] but their '
    'impact versus standard care lacks rigorous comparison. Cost-effectiveness analyses for '
    'multi-omics precision oncology are absent. Drug repurposing through Connectivity Map '
    'approaches [53] and pharmacogenomic databases [54,55] identifies candidates, and synthetic '
    'lethality -- particularly PARP inhibition in LKB1-mutant tumors [56] -- shows dual '
    'therapeutic and immunomodulatory potential. However, translation rates from computation to '
    'clinic remain low.'
)

# --- EMERGING FRONTIERS ---
doc.add_heading('Emerging frontiers', level=1)

body(
    'Single-cell RNA sequencing has revealed TME complexity invisible to bulk profiling [57,58], '
    'and spatial omics preserve tissue architecture while providing molecular detail [59,60]. '
    'Liquid biopsy technologies are approaching clinical utility: ctDNA tracks metastatic '
    'dissemination [61] and enables longitudinal risk prediction [62]. A large-scale microbiome '
    'analysis of 940 never-smoker lung cancers found no clinically relevant associations, '
    'challenging earlier positive reports [63]. Critical gaps include absent spatial omics in '
    'never-smoker NSCLC, nascent multi-modal single-cell integration, and cost barriers to '
    'clinical adoption.'
)

# --- DISCUSSION ---
doc.add_heading('Discussion and future directions', level=1)

body(
    'This review identifies interconnected challenges requiring concerted action. The feature '
    'selection philosophy divide is the most fundamental: we propose a hierarchical framework '
    'where biological clustering is performed first using unsupervised variance-based selection '
    'to reveal the natural molecular taxonomy, followed by within-subtype risk stratification '
    'integrating both molecular and clinical features. This approach preserves mechanistic '
    'insights for targeted therapy while providing clinical risk information (Fig. 1).'
)

body(
    'Feature selection methodology requires improvement through multivariate approaches (elastic '
    'net, WGCNA modules, stability selection) and functional selection tools such as MethylMix '
    'for identifying driver methylation events. Clinical variables should serve as input features, '
    'enabling formal testing of incremental multi-omics value over standard clinical predictors.'
)

body(
    'The field urgently needs dedicated multi-ethnic, multi-omics studies of never-smoker NSCLC, '
    'sex-stratified analyses as standard practice, and exposome-to-subtype mapping. AI/ML models '
    'should be developed with demographic fairness as a primary consideration. Prospective '
    'clinical validation represents the critical bottleneck: zero multi-omics studies are '
    'prospective, zero have produced companion diagnostics, and zero have informed trial enrollment.'
)

# --- CONCLUSIONS ---
doc.add_heading('Conclusions', level=1)

body(
    'Multi-omics integration and AI/ML have revealed layers of lung cancer complexity invisible '
    'to single-platform analyses. Yet translation has been limited by methodological monoculture '
    '(72% MOVICS, 88% TCGA), feature selection choices that sacrifice biological insight for '
    'prognostic convenience, and systematic exclusion of never-smokers, women, and non-European '
    'populations. Bridging these gaps through hierarchical integration, multivariate feature '
    'selection, population-specific cohorts, and prospective validation is essential to realize '
    'precision oncology for all lung cancer patients.'
)

# ============================================================
# ACKNOWLEDGMENTS
# ============================================================
doc.add_heading('Acknowledgments', level=1)
body('The authors acknowledge computational resources provided by the Digital Health Institute, '
     'Rutgers University.')

# ============================================================
# AUTHOR CONTRIBUTIONS
# ============================================================
doc.add_heading('Author contributions', level=1)
body('V.R. conceived the study, conducted the systematic literature search, performed the census '
     'of multi-omics integration studies, and wrote the manuscript. S.R. supervised the project '
     'and revised the manuscript. All authors approved the final version.')

# ============================================================
# COMPETING INTERESTS
# ============================================================
doc.add_heading('Competing interests', level=1)
body('The authors declare no competing interests.')

# ============================================================
# FIGURE LEGENDS
# ============================================================
doc.add_page_break()
doc.add_heading('Figure legends', level=1)

body(
    'Figure 1. Proposed hierarchical framework for multi-omics integration in lung cancer. '
    '(a) The current field is divided between two approaches: biological discovery (variance-based '
    'feature selection yielding 4-5 subtypes) and clinical prediction (Cox-based selection yielding '
    '2 subtypes). The prognosis-first approach collapses biologically distinct subtypes with similar '
    'survival into single categories, losing therapeutic information. (b) The proposed hierarchical '
    'framework performs biological clustering first to preserve molecular taxonomy, followed by '
    'within-subtype risk stratification using both molecular and clinical features. This approach '
    'maintains mechanistic specificity for targeted therapy while providing clinical risk information. '
    '(c) Systematic census of 25 multi-omics lung cancer studies showing MOVICS dominance (72%), '
    'TCGA dependence (88%), and the feature selection divide: 13/18 MOVICS studies use supervised '
    'Cox filtering (yielding 2 subtypes) versus 3/18 using unsupervised variance-based selection '
    '(yielding 4-5 subtypes).'
)

body(
    'Figure 2. Landscape of gaps in multi-omics lung cancer research. (a) Critical gaps identified '
    'across 12 thematic domains and 66,879 papers: sex-stratified multi-omics analyses are nearly '
    'absent, no multi-omics study has focused on never-smoker NSCLC, and the exposome-to-subtype '
    'intersection is unmapped. (b) Missing data modalities across 25 integration studies: zero '
    'metabolomics, zero spatial omics, zero DIABLO applications, minimal proteomics (2/25 studies). '
    '(c) Retraction analysis showing 76 retracted papers in the epigenetics theme, concentrated in '
    'miRNA/lncRNA prognostic signature studies. (d) Evidence strength heat map across 12 themes, '
    'ranging from strong (molecular heterogeneity, AI/ML, immune biomarkers) to weak (drug '
    'repurposing, translational outcomes).'
)

# ============================================================
# REFERENCES (max 60)
# ============================================================
doc.add_page_break()
doc.add_heading('References', level=1)

refs = [
    '1. Cancer Genome Atlas Research Network. Comprehensive molecular profiling of lung adenocarcinoma. Nature 511, 543-550 (2014).',
    '2. Cancer Genome Atlas Research Network. Comprehensive genomic characterization of squamous cell lung cancers. Nature 489, 519-525 (2012).',
    '3. Gillette, M. A. et al. Proteogenomic characterization reveals therapeutic vulnerabilities in lung adenocarcinoma. Cell 182, 200-225 (2020).',
    '4. Xu, J. Y. et al. Integrative proteomic characterization of human lung adenocarcinoma. Cell 182, 245-261 (2020).',
    '5. Chen, Y. J. et al. Proteogenomics of non-smoking lung cancer in East Asia delineates molecular signatures of pathogenesis. Cell 182, 226-244 (2020).',
    '6. Rudin, C. M. et al. Molecular subtypes of small cell lung cancer: a synthesis of human and mouse model data. Nat. Rev. Cancer 19, 289-297 (2019).',
    '7. Baine, M. K. et al. SCLC subtypes defined by ASCL1, NEUROD1, POU2F3, and YAP1. J. Thorac. Oncol. 15, 1823-1835 (2020).',
    '8. Shen, R., Olshen, A. B. & Ladanyi, M. Integrative clustering of multiple genomic data types using a joint latent variable model. Bioinformatics 25, 2906-2912 (2009).',
    '9. Wang, B. et al. Similarity network fusion for aggregating data types on a genomic scale. Nat. Methods 11, 333-337 (2014).',
    '10. Coudray, N. et al. Classification and mutation prediction from non-small cell lung cancer histopathology images using deep learning. Nat. Med. 24, 1559-1567 (2018).',
    '11. Chen, R. J. et al. Pan-cancer integrative histology-genomic analysis via multimodal deep learning. Cancer Cell 40, 865-878 (2022).',
    '12. Skoulidis, F. & Heymach, J. V. Co-occurring genomic alterations in non-small-cell lung cancer biology and therapy. Nat. Rev. Cancer 19, 495-509 (2019).',
    '13. Liu, Y. et al. Proteogenomic characterization of small cell lung cancer identifies biological insights and subtype-specific therapeutic strategies. Cell 187, 184-203 (2024).',
    '14. Martinez-Ruiz, C. et al. Genomic-transcriptomic evolution in lung cancer and metastasis. Nature 616, 553-562 (2023).',
    '15. Harada, G. et al. Rare molecular subtypes of lung cancer. Nat. Rev. Clin. Oncol. 20, 229-249 (2023).',
    '16. Argelaguet, R. et al. MOFA+: a statistical framework for comprehensive integration of multi-modal single-cell data. Genome Biol. 21, 111 (2020).',
    '17. Lu, X. et al. MOVICS: an R package for multi-omics integration and visualization in cancer subtyping. Bioinformatics 36, 5539-5541 (2020).',
    '18. Pierre-Jean, M. et al. Clustering and variable selection evaluation of 13 unsupervised methods for multi-omics data integration. Brief. Bioinform. 21, 2020-2030 (2020).',
    '19. Chauvel, C. et al. Evaluation of integrative clustering methods for the analysis of multi-omics data. Brief. Bioinform. 21, 541-552 (2020).',
    '20. Athaya, T. et al. Multimodal deep learning approaches for single-cell multi-omics data integration. Brief. Bioinform. 24, bbad313 (2023).',
    '21. Hu, B. et al. Benchmarking algorithms for single-cell multi-omics prediction and integration. Nat. Methods 21, 2126-2137 (2024).',
    '22. Baiao, A. et al. A technical review of multi-omics data integration methods. Brief. Bioinform. 26, bbaf066 (2025).',
    '23. Zou, Y. et al. Multi-omics consensus portfolio to refine the classification of lung adenocarcinoma. Transl. Lung Cancer Res. 11, 2251-2271 (2022).',
    '24. Ravichandran, V. Multi-omics integration for molecular subtyping in lung cancer: a comprehensive review. Digital Health Institute, Rutgers University (2026).',
    '25. Wang, X. et al. A pathology foundation model for cancer diagnosis and prognosis prediction. Nature 634, 921-928 (2024).',
    '26. Yang, G. et al. A foundation model for generalizable cancer diagnosis and survival prediction. Nat. Commun. 16, 2045 (2025).',
    '27. Rakaee, M. et al. Deep learning model for predicting immunotherapy response in advanced NSCLC. JAMA Oncol. 11, 45-54 (2025).',
    '28. Arshad, M. et al. Artificial intelligence and machine learning in lung cancer. Cancers 17, 456 (2025).',
    '29. Siegfried, J. M. et al. Women and lung cancer: does oestrogen play a role? Lancet Oncol. 2, 506-513 (2001).',
    '30. Park, E. et al. Proteogenomic characterization reveals estrogen signaling as a target for never-smoker LUAD. Cancer Res. 84, 2178-2192 (2024).',
    '31. Conforti, F. et al. Cancer immunotherapy efficacy and patients\' sex: a systematic review and meta-analysis. Lancet Oncol. 19, 737-746 (2018).',
    '32. Vavala, T. et al. Gender differences and immunotherapy outcome in advanced lung cancer. Int. J. Mol. Sci. 22, 11888 (2021).',
    '33. Zhang, T. et al. Genomic and evolutionary classification of lung cancer in never smokers. Nat. Genet. 53, 1348-1359 (2021).',
    '34. Hamouz, R. Z. et al. A functional genomics review of non-small-cell lung cancer in never smokers. Int. J. Mol. Sci. 24, 13314 (2023).',
    '35. Pershagen, G. et al. Residential radon exposure and lung cancer in Sweden. N. Engl. J. Med. 330, 159-164 (1994).',
    '36. Vahakangas, K. H. et al. Mutations of p53 and ras genes in radon-associated lung cancer. Lancet 339, 576-580 (1992).',
    '37. Wielscher, M. et al. DNA methylation signature of chronic low-grade inflammation and its role in cardio-respiratory disease. Nat. Commun. 13, 2536 (2022).',
    '38. Alam, H. et al. KMT2D deficiency impairs super-enhancers to confer a glycolytic vulnerability in lung cancer. Cancer Cell 37, 599-617 (2020).',
    '39. de Miguel, F. J. et al. Mammalian SWI/SNF chromatin remodeling complexes promote TKI resistance in EGFR-mutant lung cancer. Cancer Cell 41, 1543-1560 (2023).',
    '40. Yuan, G. et al. Elevated NSD3 histone methylation activity drives squamous cell lung cancer. Nature 590, 504-508 (2021).',
    '41. Heeke, S. et al. Tumor- and circulating-free DNA methylation identifies clinically relevant SCLC subtypes. Cancer Cell 42, 225-242 (2024).',
    '42. Zuccato, C. F. et al. Prediction of brain metastasis development with DNA methylation signatures. Nat. Med. 31, 456-467 (2025).',
    '43. Rizvi, N. A. et al. Mutational landscape determines sensitivity to PD-1 blockade in NSCLC. Science 348, 124-128 (2015).',
    '44. Thorsson, V. et al. The immune landscape of cancer. Immunity 48, 812-830 (2018).',
    '45. Charoentong, P. et al. Pan-cancer immunogenomic analyses reveal genotype-immunophenotype relationships. Cell Rep. 18, 248-262 (2017).',
    '46. Bagaev, A. et al. Conserved pan-cancer microenvironment subtypes predict response to immunotherapy. Cancer Cell 39, 845-865 (2021).',
    '47. Stein-Thoeringer, C. K. et al. A non-antibiotic-disrupted gut microbiome is associated with clinical responses to immunotherapy. Nat. Med. 29, 906-916 (2023).',
    '48. Ferrari, V. et al. Sensitizing cancer cells to immune checkpoint inhibitors by microbiota-mediated upregulation of HLA class I. Cancer Cell 41, 1717-1730 (2023).',
    '49. Bach, P. B. et al. Racial differences in the treatment of early-stage lung cancer. N. Engl. J. Med. 341, 1198-1205 (1999).',
    '50. Dutta, R. et al. Understanding inequities in precision oncology diagnostics. Nat. Cancer 4, 787-794 (2023).',
    '51. Tsimberidou, A. M. et al. Molecular tumour boards -- current and future considerations for precision oncology. Nat. Rev. Clin. Oncol. 20, 843-863 (2023).',
    '52. Kato, S. et al. Real-world data from a molecular tumor board demonstrates improved outcomes. Nat. Commun. 11, 4965 (2020).',
    '53. Lamb, J. et al. The Connectivity Map: a new tool for biomedical research. Nat. Rev. Cancer 7, 54-60 (2007).',
    '54. Cannon, M. et al. DGIdb 5.0: rebuilding the drug-gene interaction database. Nucleic Acids Res. 52, D1227-D1235 (2024).',
    '55. Nair, N. U. et al. A landscape of response to drug combinations in NSCLC. Nat. Commun. 14, 3642 (2023).',
    '56. Long, L. et al. PARP inhibition induces synthetic lethality and adaptive immunity in LKB1-mutant lung cancer. Cancer Res. 83, 568-581 (2023).',
    '57. Lambrechts, D. et al. Phenotype molding of stromal cells in the lung tumor microenvironment. Nat. Med. 24, 1277-1289 (2018).',
    '58. Maynard, A. et al. Therapy-induced evolution of human lung cancer revealed by single-cell RNA sequencing. Cell 182, 1232-1251 (2020).',
    '59. Tagore, S. et al. Single-cell and spatial genomic landscape of NSCLC brain metastases. Nat. Med. 31, 678-690 (2025).',
    '60. Li, Z. et al. AI-enabled virtual spatial proteomics from histopathology. Nat. Med. 32, 234-246 (2026).',
]

for r in refs:
    ref_entry(r)

# ============================================================
# SAVE
# ============================================================
output = '/Users/vravichandran/Documents/Rutgers_DHI/LungCancer_Review/Ravichandran_MultiOmics_LungCancer_npjPO_2026.docx'
doc.save(output)

# Word count
import re
text = []
for para in doc.paragraphs:
    if para.style.name.startswith('Heading') or para.style.name == 'ManuscriptTitle':
        continue
    text.append(para.text)
full = ' '.join(text)
words = len(re.findall(r'\b\w+\b', full))
print(f'Manuscript saved: {output}')
print(f'References: {len(refs)}')
print(f'Approximate word count (all text): {words}')
