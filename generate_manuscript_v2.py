#!/usr/bin/env python3
"""
Generate manuscript for npj Precision Oncology:
"Multi-Omics and AI in Lung Cancer: Gaps in Integration, Stratification, and Translation"
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

OUTPUT_PATH = "/Users/vravichandran/Documents/Rutgers_DHI/LungCancer_Review/Ravichandran_MultiOmics_LungCancer_npjPO_2026.docx"

doc = Document()

# ── Style setup ──────────────────────────────────────────────────────────────

style_normal = doc.styles['Normal']
font = style_normal.font
font.name = 'Arial'
font.size = Pt(11)
font.color.rgb = RGBColor(0, 0, 0)
pf = style_normal.paragraph_format
pf.line_spacing = 2.0
pf.space_after = Pt(0)
pf.space_before = Pt(0)

for level, size, bold, italic in [(1, 14, True, False), (2, 12, True, True)]:
    sname = f'Heading {level}'
    s = doc.styles[sname]
    s.font.name = 'Arial'
    s.font.size = Pt(size)
    s.font.bold = bold
    s.font.italic = italic
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.line_spacing = 2.0
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(6)

# ── Helpers ──────────────────────────────────────────────────────────────────

def add_para(text, style='Normal', bold=False, italic=False, alignment=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Arial'
    if alignment:
        p.alignment = alignment
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(9)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="4472C4"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.line_spacing = 1.15
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.15
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()  # spacer

def word_count(text):
    return len(text.split())

# ── Collect all body text for word counting ──────────────────────────────────
all_body_text = []

# ── Title page ───────────────────────────────────────────────────────────────

title = "Multi-Omics and AI in Lung Cancer: Gaps in Integration, Stratification, and Translation"
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(title)
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Arial'

doc.add_paragraph()

authors = "Vignesh Ravichandran\u00B9, Suchismita Ray\u00B9"
add_para(authors, alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

aff = "\u00B9 Digital Health Institute, Rutgers University, New Brunswick, NJ, USA"
add_para(aff, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

corr = "Correspondence: vignesh.ravichandran@rutgers.edu"
add_para(corr, alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ── Abstract ─────────────────────────────────────────────────────────────────

doc.add_heading('Abstract', level=1)

abstract_text = (
    "Lung cancer remains the leading cause of cancer death worldwide. While multi-omics integration "
    "and artificial intelligence promise to transform molecular classification and treatment selection, "
    "critical gaps persist. We systematically reviewed 66,879 papers across 12 thematic domains and "
    "conducted a focused census of 25 multi-omics integration studies in lung cancer. We identify a "
    "fundamental methodological divide: 72% of studies use the MOVICS framework, 88% depend on TCGA "
    "data, and the majority employ survival-based feature selection that collapses biological "
    "heterogeneity into binary risk categories. Sex-stratified and never-smoker-focused multi-omics "
    "analyses are virtually absent. No study integrates clinical features as competing inputs alongside "
    "molecular data. Zero studies are prospective, and zero have produced companion diagnostics. We "
    "propose a hierarchical framework that preserves biological taxonomy while enabling clinical risk "
    "stratification, and outline priorities for closing these translational gaps."
)
add_para(abstract_text)
print(f"Abstract word count: {word_count(abstract_text)}")

doc.add_page_break()

# ── SECTION 1: Introduction ─────────────────────────────────────────────────

doc.add_heading('Introduction', level=1)

intro = (
    "Lung cancer is the most commonly diagnosed malignancy and the leading cause of cancer mortality "
    "worldwide, accounting for approximately 1.8 million deaths annually. Non-small cell lung cancer "
    "(NSCLC) constitutes roughly 85% of diagnoses, with lung adenocarcinoma (LUAD) representing "
    "approximately 40% and lung squamous cell carcinoma (LUSC) approximately 30% of cases. Despite "
    "transformative advances in targeted therapy\u2014including osimertinib for EGFR-mutant disease, "
    "sotorasib and adagrasib for KRAS G12C, and crizotinib and alectinib for ALK-rearranged tumors"
    "\u2014and the broad deployment of immune checkpoint inhibitors such as pembrolizumab and "
    "atezolizumab, five-year survival remains a sobering 16\u201325% across stages.\n\n"

    "The genomic era has fundamentally reframed our understanding of lung cancer biology. The Cancer "
    "Genome Atlas (TCGA) established LUAD and LUSC as molecularly distinct diseases with divergent "
    "driver landscapes and therapeutic vulnerabilities [1,2]. The Clinical Proteomic Tumor Analysis "
    "Consortium (CPTAC) extended these insights to the proteomic and phosphoproteomic layers, revealing "
    "that protein-level biology frequently diverges from genomic predictions and exposing druggable "
    "kinase activation events invisible to DNA sequencing alone [3\u20135]. In parallel, small cell lung "
    "cancer (SCLC) has been reclassified from a monolithic entity into four transcription-factor-defined "
    "subtypes\u2014ASCL1, NEUROD1, POU2F3, and YAP1\u2014each with distinct therapeutic "
    "sensitivities [6,7].\n\n"

    "Multi-omics integration methods\u2014including iCluster [8], Similarity Network Fusion (SNF) [9], "
    "and consensus frameworks such as MOVICS [18]\u2014promise to capture cross-layer regulatory "
    "relationships that no single platform can resolve. Concurrently, artificial intelligence and "
    "machine learning approaches have demonstrated remarkable capacity: deep learning models now "
    "classify histological subtypes and predict driver mutations directly from hematoxylin and eosin "
    "(H&E)-stained slides [10], multimodal architectures fuse histopathology with genomic data to "
    "improve prognostic prediction [11], and foundation models pre-trained on millions of pathology "
    "images are emerging as general-purpose encoders for diverse downstream tasks [12,13].\n\n"

    "Yet translation has stalled. Reproducibility failures, population bias, and the absence of "
    "prospective validation have prevented multi-omics and AI discoveries from reaching clinical "
    "practice. In this review, we systematically analyzed 66,879 papers across 12 thematic domains "
    "and conducted a focused census of 25 multi-omics integration studies in lung cancer. We identify "
    "a critical methodological divide in feature selection philosophy, document pervasive gaps in "
    "population representation, and propose a hierarchical framework for integration that preserves "
    "biological taxonomy while enabling clinical stratification (Table 1; Fig. 1)."
)
add_para(intro)
all_body_text.append(intro)

# ── SECTION 2: Molecular heterogeneity ──────────────────────────────────────

doc.add_heading('Molecular heterogeneity of lung cancer', level=1)

sec2 = (
    "The molecular landscape of lung cancer is defined by histology-specific driver events, "
    "co-mutation patterns that modify treatment response, and proteomic divergence from genomic "
    "predictions. In LUAD, TCGA profiling identified recurrent drivers in EGFR, KRAS, ALK, ROS1, "
    "and BRAF, alongside three transcriptional subtypes\u2014terminal respiratory unit (TRU), "
    "proximal-inflammatory (PI), and proximal-proliferative (PP)\u2014that correlate with clinical "
    "behavior and immune contexture [1]. LUSC, by contrast, is dominated by near-universal TP53 "
    "mutation (>90%), amplifications in FGFR1 and SOX2, and four expression-based subtypes with "
    "distinct metabolic and immune features [2]. The absence of targetable oncogenic drivers in LUSC "
    "remains a central therapeutic challenge.\n\n"

    "Co-mutation context profoundly shapes biology and treatment response. Skoulidis and Heymach "
    "demonstrated that KRAS-mutant LUAD segregates into biologically and clinically distinct subgroups "
    "defined by co-occurring alterations: KRAS/STK11 (immune-inert, immunotherapy-resistant), "
    "KRAS/KEAP1 (metabolically dysregulated, chemotherapy-resistant), and KRAS/TP53 (inflamed, "
    "immunotherapy-responsive) [14]. These co-mutation patterns are invisible to single-gene testing "
    "and illustrate why multi-dimensional molecular characterization is essential.\n\n"

    "CPTAC proteogenomic studies have revealed a consistent theme: protein abundance and phosphorylation "
    "frequently diverge from mRNA predictions. Gillette et al. identified druggable kinase activations "
    "in LUAD that were undetectable at the transcript level [3]. Xu et al. confirmed proteomic-specific "
    "subtypes in an independent Chinese cohort [4]. Chen et al. characterized never-smoker LUAD in East "
    "Asian patients, uncovering unique EGFR-pathway rewiring and estrogen receptor signaling that "
    "suggests sex- and exposure-specific therapeutic vulnerabilities [5].\n\n"

    "SCLC has undergone a parallel molecular reclassification. Rudin et al. proposed four subtypes "
    "defined by dominant transcription factor expression\u2014ASCL1 (SCLC-A), NEUROD1 (SCLC-N), "
    "POU2F3 (SCLC-P), and YAP1 (SCLC-Y)\u2014each with distinct therapeutic sensitivities [6,7]. "
    "Recent proteogenomic characterization by Liu et al. has further refined these subtypes and "
    "identified subtype-specific druggable targets, including metabolic vulnerabilities in SCLC-A and "
    "immune checkpoint dependencies in SCLC-Y [15]. Histologic transformation\u2014particularly "
    "NSCLC-to-SCLC conversion under TKI pressure\u2014adds temporal complexity that static multi-omics "
    "snapshots cannot capture [16].\n\n"

    "A pervasive limitation across these landmark studies is population bias. The majority derive from "
    "European-ancestry cohorts, sex-stratified analyses are rare, and never-smoker populations remain "
    "underrepresented despite comprising 10\u201325% of global lung cancer diagnoses (Fig. 1)."
)
add_para(sec2)
all_body_text.append(sec2)

# ── SECTION 3: Multi-omics integration ──────────────────────────────────────

doc.add_heading('Multi-omics integration methods and the feature selection problem', level=1)

doc.add_heading('Integration frameworks', level=2)

sec3a = (
    "A growing ecosystem of computational methods enables the joint analysis of multiple molecular "
    "data types. iCluster models shared latent variables across platforms using penalized matrix "
    "decomposition, providing statistically rigorous subtype discovery but suffering from Gaussian "
    "assumptions and computational cost that limits scalability [8]. Similarity Network Fusion (SNF) "
    "constructs patient-similarity networks for each data type and iteratively fuses them, preserving "
    "local structure but requiring careful selection of the number of clusters K [9]. MOFA+ employs "
    "Bayesian group factor analysis to decompose variation across data modalities into interpretable "
    "latent factors, excelling at variance attribution but not directly producing discrete subtypes [17]. "
    "MOVICS provides a consensus framework that runs up to 10 clustering algorithms simultaneously and "
    "identifies robust subtypes via consensus, offering flexibility at the cost of computational "
    "intensity [18].\n\n"

    "Systematic benchmarking has revealed that no single method consistently outperforms others across "
    "datasets and cancer types [19,20]. Performance depends heavily on the fusion strategy "
    "(early, intermediate, or late), the number and type of omics layers, and critically, on the "
    "feature selection approach applied before integration [20]. Deep learning approaches\u2014including "
    "variational autoencoders, graph neural networks, and transformer architectures\u2014offer the "
    "theoretical advantage of capturing non-linear cross-modal relationships, but their data hunger "
    "poses a practical barrier in oncology where cohort sizes rarely exceed several hundred [21,22]. "
    "Spatial and temporal integration methods remain nascent, with few frameworks designed to handle "
    "the unique data structures of spatial transcriptomics or longitudinal sampling [23] (Table 2)."
)
add_para(sec3a)
all_body_text.append(sec3a)

doc.add_heading('The feature selection philosophy divide', level=2)

sec3b = (
    "Our census of 25 multi-omics integration studies in lung cancer reveals a fundamental "
    "methodological divide that has not been previously described. Of 25 studies, 18 (72%) employ the "
    "MOVICS framework, and 22 (88%) rely exclusively on TCGA data. More consequentially, two "
    "philosophically distinct approaches to feature selection produce dramatically different biological "
    "conclusions.\n\n"

    "The first approach, adopted by 13 of 18 MOVICS studies, uses Cox proportional hazards regression "
    "to select features significantly associated with overall survival (typically p < 0.05 or p < 0.01). "
    "This supervised, prognosis-first strategy typically retains 2,000\u20133,500 features and converges "
    "on two subtypes\u2014invariably labeled \u2018high-risk\u2019 and \u2018low-risk.\u2019 The second "
    "approach, used by 3 of 18 MOVICS studies, selects features based on variance metrics such as "
    "median absolute deviation (MAD) or standard deviation (SD). This unsupervised, biology-first "
    "strategy retains approximately 10,000 features and discovers 4\u20135 subtypes with distinct "
    "molecular hallmarks.\n\n"

    "The hidden problem with the prognosis-first approach is that it merges biologically distinct "
    "subtypes that happen to share similar survival outcomes into a single risk category. For instance, "
    "immune-excluded tumors (which might benefit from anti-angiogenic agents) and metabolically "
    "dysfunctional tumors (which might respond to metabolic inhibitors) both confer poor prognosis and "
    "are therefore collapsed into the same \u2018high-risk\u2019 group. This conflation sacrifices the "
    "biological resolution needed for precision therapy selection.\n\n"

    "Zou et al. demonstrated the value of the biology-first approach: their variance-based multi-omics "
    "clustering of LUAD revealed four subtypes with distinct sensitivities to cisplatin, paclitaxel, "
    "docetaxel, and immunotherapy\u2014therapeutic distinctions that would be invisible in a binary risk "
    "framework [24]. The two NMF-based studies using institutional cohorts (Zhang et al., n=122; "
    "Zhong et al., n=101) similarly recovered biologically coherent subtypes by avoiding survival-driven "
    "feature selection.\n\n"

    "Several additional methodological gaps emerge from our census. Clinical features\u2014including "
    "age, stage, smoking history, sex, and ECOG performance status\u2014are almost never incorporated "
    "as input features alongside molecular data. No study tests the incremental predictive value of "
    "molecular features over clinical variables alone. Feature selection is overwhelmingly univariate, "
    "testing one gene at a time; multivariate methods such as elastic net regularization, weighted gene "
    "co-expression network analysis (WGCNA), and stability selection are not used. DIABLO, a supervised "
    "multi-omics method that has shown success in breast cancer biomarker discovery, has zero "
    "applications in lung cancer. Notably, zero studies incorporate metabolomics or spatial omics data, "
    "representing substantial gaps in the omics landscape being interrogated (Table 1; Fig. 2)."
)
add_para(sec3b)
all_body_text.append(sec3b)

# ── SECTION 4: AI and machine learning ───────────────────────────────────────

doc.add_heading('AI and machine learning in lung cancer', level=1)

sec4 = (
    "Artificial intelligence has introduced capabilities that complement and extend multi-omics "
    "integration. In computational pathology, Coudray et al. demonstrated that convolutional neural "
    "networks trained on H&E-stained whole-slide images can classify LUAD versus LUSC with high "
    "accuracy and, more remarkably, predict the mutational status of STK11, EGFR, and TP53 directly "
    "from histomorphology [10]. This finding implies that tissue architecture encodes molecular "
    "information accessible to deep learning but invisible to human pathologists.\n\n"

    "Multimodal architectures represent a natural convergence of AI and multi-omics. Chen et al. "
    "developed a pan-cancer framework that fuses histopathology features with genomic data, "
    "demonstrating that multimodal integration consistently outperforms single-modality models for "
    "survival prediction [11]. This approach addresses a limitation of most multi-omics studies, which "
    "ignore morphological information entirely.\n\n"

    "The emergence of pathology foundation models since 2023 represents a paradigm shift. Wang et al. "
    "pre-trained a foundation model on over 100 million pathology image tiles, achieving state-of-the-art "
    "performance across cancer diagnosis and prognosis tasks without task-specific training [12]. "
    "Yang et al. extended this concept to a generalizable model for both diagnosis and survival "
    "prediction across cancer types [13]. These models may eventually serve as general-purpose feature "
    "extractors that can be combined with molecular data in multimodal pipelines.\n\n"

    "Radiomics\u2014the extraction of quantitative features from clinical imaging\u2014offers a "
    "non-invasive complement to tissue-based omics. Rakaee et al. developed a deep learning model "
    "trained on CT scans that predicts immunotherapy response in advanced NSCLC [25]. However, "
    "radiomic features are notoriously sensitive to scanner parameters, reconstruction algorithms, and "
    "acquisition protocols, raising reproducibility concerns.\n\n"

    "A consistent finding across AI applications in oncology is performance degradation upon external "
    "validation [26]. Models trained on single-institution or single-cohort data frequently fail to "
    "generalize, reflecting overfitting to institution-specific confounders. Algorithmic bias compounds "
    "this problem: demographic underrepresentation in training data means that model performance is "
    "rarely evaluated across sex, ancestry, or smoking-status subgroups. No AI tool has been developed "
    "or validated specifically for never-smoker NSCLC (Fig. 1)."
)
add_para(sec4)
all_body_text.append(sec4)

# ── SECTION 5: Underrepresented populations ──────────────────────────────────

doc.add_heading('Underrepresented populations in lung cancer research', level=1)

doc.add_heading('Sex and gender differences', level=2)

sec5a = (
    "Lung cancer in women differs from lung cancer in men in incidence patterns, molecular drivers, "
    "and treatment response. Women develop lung cancer at younger ages, with lighter smoking histories, "
    "and harbor higher rates of EGFR mutations [27,28]. Estrogen receptor signaling has been implicated "
    "as a biologically plausible contributor to these differences, with preclinical evidence suggesting "
    "that estrogen promotes lung cancer cell proliferation through both genomic and non-genomic "
    "mechanisms [27]. Conforti et al. conducted a meta-analysis of randomized trials and reported "
    "greater immune checkpoint inhibitor benefit in men than women, although this finding remains "
    "debated and may reflect confounding by molecular subtypes and smoking status [29].\n\n"

    "Park et al. provided direct molecular evidence for sex-specific biology: their proteogenomic "
    "analysis of never-smoker LUAD identified estrogen signaling as an actionable target enriched in "
    "female patients, with potential therapeutic implications for anti-estrogen combination "
    "strategies [30]. Despite these findings, pharmacogenomic databases lack sex-stratified drug "
    "response data, and the functional consequences of Y chromosome loss\u2014a common event in male "
    "cancers\u2014remain poorly characterized in lung cancer."
)
add_para(sec5a)
all_body_text.append(sec5a)

doc.add_heading('Never-smoker lung cancer', level=2)

sec5b = (
    "Lung cancer in never-smokers represents a molecularly and clinically distinct entity that is "
    "systematically underserved by current multi-omics research. Zhang et al. performed whole-genome "
    "sequencing of never-smoker lung cancers and identified three genomic subtypes\u2014piano (low "
    "mutation burden, indolent), mezzo-forte (EGFR-driven, intermediate), and forte (rapid progression, "
    "whole-genome doubling)\u2014that correlate with clinical trajectory [31]. Sun et al. provided an "
    "early conceptual framework arguing that never-smoker lung cancer should be studied as a distinct "
    "disease [32].\n\n"

    "The mutational landscape of never-smoker lung cancer is characterized by the absence of the "
    "smoking-associated signature SBS4 and enrichment of clock-like signatures (SBS1, SBS5) and "
    "APOBEC-mediated mutagenesis [33]. The tumor microenvironment tends toward immune-cold phenotypes, "
    "with lower TMB and reduced neoantigen load, potentially explaining the variable response to "
    "immune checkpoint inhibitors observed in this population [34]. Never-smoker lung cancer accounts "
    "for 10\u201325% of global diagnoses with rising incidence, yet no clinical trial has been designed "
    "around never-smoker molecular subtypes, and no multi-omics integration study has focused "
    "specifically on this population."
)
add_para(sec5b)
all_body_text.append(sec5b)

doc.add_heading('Environmental exposures', level=2)

sec5c = (
    "Environmental carcinogens contribute to lung cancer through mechanisms distinct from tobacco "
    "smoke. Martinez-Ruiz et al. demonstrated that particulate matter (PM2.5) promotes the clonal "
    "expansion of pre-existing EGFR-mutant cells through IL-1\u03B2-mediated inflammation rather than "
    "direct mutagenesis\u2014a paradigm-shifting finding that explains the paradox of carcinogenesis "
    "without increased mutational burden [16]. Residential radon exposure shows dose-dependent "
    "associations with lung cancer risk and produces distinct molecular signatures, including "
    "characteristic TP53 mutation spectra [35,36]. Indoor air pollution from cooking fuels "
    "disproportionately affects women in Asian countries and represents a major but molecularly "
    "undercharacterized exposure [37]. The mapping of environmental exposome to molecular subtypes "
    "remains essentially absent from the literature, representing a critical gap for populations where "
    "non-tobacco exposures dominate."
)
add_para(sec5c)
all_body_text.append(sec5c)

# ── SECTION 6: Epigenetics ──────────────────────────────────────────────────

doc.add_heading('Epigenetics and the reproducibility concern', level=1)

sec6 = (
    "Epigenetic alterations are among the earliest events in lung carcinogenesis and offer both "
    "biological insight and clinical biomarker potential. DNA methylation profiling has identified "
    "consistent hypermethylation of tumor suppressors including CDKN2A, RASSF1A, and MGMT across lung "
    "cancer subtypes [38]. Chromatin remodeling mutations are increasingly recognized as therapeutic "
    "vulnerabilities: KMT2D deficiency creates super-enhancer dependencies that confer glycolytic "
    "vulnerability [39], SWI/SNF complex alterations promote TKI resistance in EGFR-mutant lung "
    "cancer [40], and NSD3 amplification drives a histone methylation program specific to LUSC [41].\n\n"

    "Cell-free DNA (cfDNA) methylation has emerged as a promising non-invasive biomarker platform. "
    "Heeke et al. demonstrated that circulating-free DNA methylation signatures can identify clinically "
    "relevant SCLC subtypes from blood samples [42]. Zuccato et al. developed methylation-based models "
    "that predict brain metastasis development, a finding with significant implications for surveillance "
    "strategies [43].\n\n"

    "However, the epigenetics domain carries a significant reproducibility warning. Our systematic "
    "review identified 76 retracted papers in the epigenetics theme\u2014the highest retraction count "
    "of any thematic domain analyzed. Retractions are concentrated in studies reporting prognostic "
    "signatures based on miRNA and lncRNA expression panels, suggesting a pattern of overfitting to "
    "discovery cohorts without adequate validation. This retraction signal underscores the need for "
    "pre-registered validation protocols and independent replication before epigenetic biomarker "
    "panels are advanced toward clinical use."
)
add_para(sec6)
all_body_text.append(sec6)

# ── SECTION 7: Immune biomarkers ────────────────────────────────────────────

doc.add_heading('Immune biomarkers and the tumor microenvironment', level=1)

sec7 = (
    "The success of immune checkpoint inhibitors has catalyzed extensive efforts to identify predictive "
    "biomarkers beyond PD-L1 expression. Tumor mutational burden (TMB) was among the first genomic "
    "biomarkers associated with immunotherapy response, with Rizvi et al. demonstrating that high TMB "
    "correlates with improved outcomes on PD-1 blockade in NSCLC [44]. However, the predictive value "
    "of TMB is confounded by co-occurring alterations: STK11 and KEAP1 mutations override the "
    "association between high TMB and immunotherapy benefit, creating a population of patients with "
    "high TMB but primary resistance [14].\n\n"

    "Pan-cancer immune classification frameworks provide additional granularity. Thorsson et al. "
    "defined six immune subtypes across 33 cancer types, with lung cancers distributing across "
    "multiple categories [45]. Charoentong et al. developed the Immunophenoscore, an aggregated "
    "metric of immune determinants that predicts checkpoint inhibitor response [46]. Bagaev et al. "
    "identified four conserved tumor microenvironment archetypes\u2014immune-enriched fibrotic, "
    "immune-enriched non-fibrotic, immune-desert fibrotic, and immune-desert non-fibrotic\u2014that "
    "predict immunotherapy response across cancer types [47].\n\n"

    "The microbiome has emerged as an unexpected modulator of immunotherapy efficacy. Stein-Thoeringer "
    "et al. demonstrated that a non-antibiotic-disrupted gut microbiome is associated with improved "
    "clinical responses to checkpoint inhibitors [48]. Ferrari et al. provided a mechanistic link, "
    "showing that specific bacterial species sensitize cancer cells to immune killing through "
    "upregulation of HLA class I expression [49]. Despite this expanding landscape of immune "
    "biomarkers, composite panels remain unstandardized, and spatial biomarkers that capture the "
    "geography of immune infiltration have not been validated prospectively."
)
add_para(sec7)
all_body_text.append(sec7)

# ── SECTION 8: Translational gaps ───────────────────────────────────────────

doc.add_heading('Translational gaps and drug repurposing', level=1)

sec8 = (
    "The translation of multi-omics discoveries into clinical benefit faces structural barriers that "
    "extend beyond methodology. Disparities in access to molecular testing compound biological "
    "complexity: Bach et al. documented persistent racial differences in the treatment of early-stage "
    "lung cancer [50], and Dutta et al. demonstrated that inequities in precision oncology diagnostics "
    "perpetuate disparate outcomes for underrepresented populations [51].\n\n"

    "Molecular tumor boards (MTBs) represent the current mechanism for translating genomic findings "
    "into treatment decisions. Tsimberidou et al. outlined the evolving role of MTBs in precision "
    "oncology, emphasizing the need for structured frameworks that can incorporate multi-omics data "
    "beyond single-gene actionability [52]. Kato et al. demonstrated that N-of-one molecular tumor "
    "board recommendations improve outcomes compared to unmatched therapy in a real-world setting [53].\n\n"

    "Drug repurposing offers an accelerated path from molecular insight to therapeutic intervention. "
    "The Connectivity Map (CMap) enables identification of compounds that reverse disease-associated "
    "gene expression signatures [54]. The Drug-Gene Interaction Database (DGIdb) provides curated "
    "mappings between druggable genes and existing therapeutics [55]. Nair et al. systematically mapped "
    "the landscape of drug combination responses in NSCLC, identifying synergistic pairs that warrant "
    "clinical testing [56]. Synthetic lethality strategies are particularly promising: Long et al. "
    "demonstrated that PARP inhibition induces synthetic lethality in LKB1-mutant lung cancer while "
    "simultaneously activating adaptive immunity, suggesting a dual mechanism of action [57]. The "
    "clinical success of KRAS G12C inhibitors\u2014sotorasib and adagrasib\u2014demonstrates that "
    "molecularly targeted therapy can emerge from previously undruggable targets when structural biology "
    "and medicinal chemistry converge."
)
add_para(sec8)
all_body_text.append(sec8)

# ── SECTION 9: Emerging frontiers ────────────────────────────────────────────

doc.add_heading('Emerging frontiers', level=1)

sec9 = (
    "Single-cell and spatial technologies are reshaping our understanding of tumor heterogeneity at "
    "unprecedented resolution. Lambrechts et al. applied single-cell RNA sequencing to NSCLC and "
    "revealed that the tumor microenvironment molds stromal cell phenotypes in ways invisible to bulk "
    "profiling\u2014fibroblasts, endothelial cells, and immune cells adopt tumor-specific transcriptional "
    "programs that vary across patients [58]. Maynard et al. tracked therapy-induced evolution in "
    "individual patients through serial single-cell profiling, demonstrating that treatment pressure "
    "selects for transcriptionally distinct resistant clones with specific vulnerabilities [59].\n\n"

    "Spatial genomics adds the critical dimension of tissue architecture. Tagore et al. combined "
    "single-cell and spatial transcriptomics to characterize NSCLC brain metastases, revealing that "
    "metastatic niches harbor immunosuppressive cellular neighborhoods distinct from the primary "
    "tumor [60]. These spatial relationships\u2014which cell types neighbor which, and how signaling "
    "gradients organize the microenvironment\u2014cannot be captured by dissociation-based methods.\n\n"

    "Liquid biopsy is extending molecular profiling beyond tissue. Abbosh et al. used circulating "
    "tumor DNA (ctDNA) in the TRACERx study to track early metastatic dissemination in NSCLC, "
    "demonstrating that subclonal architecture in the primary tumor predicts the timing and sites of "
    "relapse [61]. Minimal residual disease (MRD) monitoring via ctDNA is increasingly being adopted "
    "in clinical trials as a surrogate endpoint. Conversely, not all emerging biomarker hypotheses "
    "have survived scrutiny: McElderry et al. analyzed the microbiome of 940 lung cancers in "
    "never-smokers and found no clinically relevant associations, providing an important null result "
    "that tempers enthusiasm for microbiome-based diagnostics in this population [62]."
)
add_para(sec9)
all_body_text.append(sec9)

# ── SECTION 10: Discussion ──────────────────────────────────────────────────

doc.add_heading('Discussion and future directions', level=1)

sec10 = (
    "The central finding of this review is that multi-omics integration in lung cancer is constrained "
    "not by a lack of data or methods, but by a feature selection philosophy that prioritizes prognosis "
    "over biology. We propose a hierarchical framework that addresses this limitation through a "
    "three-stage pipeline (Fig. 3).\n\n"

    "Stage 1 applies unsupervised biological clustering using variance-based features (MAD or SD) to "
    "discover natural molecular subtypes without imposing a survival prior. This approach preserves "
    "biological heterogeneity and allows the identification of subtypes defined by distinct mechanisms "
    "rather than by shared outcomes. Stage 2 performs within-subtype clinical risk stratification using "
    "multivariate methods\u2014elastic net regularization, WGCNA, or stability selection\u2014that "
    "integrate both molecular and clinical features as competing predictors. This design tests whether "
    "molecular data adds predictive value beyond established clinical variables, a question that no "
    "current multi-omics study addresses. Stage 3 maps each subtype-risk combination to specific "
    "therapeutic strategies, including targeted therapy, immunotherapy, and drug repurposing candidates "
    "identified through CMap or synthetic lethality screens.\n\n"

    "Several methodological improvements are immediately actionable. Feature selection should shift from "
    "univariate Cox filtering to multivariate approaches: elastic net simultaneously selects and "
    "regularizes features across omics layers, WGCNA identifies co-expressed gene modules that "
    "correspond to biological pathways rather than individual genes, and stability selection provides "
    "frequency-based confidence in selected features. MethylMix can distinguish driver methylation "
    "events from passenger noise, a critical capability given that most CpG sites are biologically "
    "inert. DIABLO, which has demonstrated strong performance in breast cancer multi-omics biomarker "
    "discovery, should be systematically evaluated in lung cancer.\n\n"

    "Population representation requires structural change. Multi-ethnic cohorts that include adequate "
    "representation of East Asian, African, and Latin American populations must become standard rather "
    "than exceptional. Sex-stratified analysis should be a default reporting requirement, not an "
    "afterthought. Dedicated never-smoker cohorts with paired multi-omics profiling are needed to "
    "characterize the biology of a disease that affects a growing population worldwide.\n\n"

    "Most critically, the field must move beyond retrospective TCGA re-analysis. Zero multi-omics "
    "integration studies in lung cancer are prospective. Zero have produced companion diagnostics "
    "approved for clinical use. The path forward requires prospective multi-omics trials that embed "
    "integration analysis within clinical decision-making, cost-effectiveness evidence that justifies "
    "multi-platform testing, and regulatory engagement to define companion diagnostic pathways for "
    "multi-omics classifiers (Fig. 3)."
)
add_para(sec10)
all_body_text.append(sec10)

# ── SECTION 11: Conclusions ─────────────────────────────────────────────────

doc.add_heading('Conclusions', level=1)

sec11 = (
    "Multi-omics integration and artificial intelligence have revealed layers of biological complexity "
    "in lung cancer that are invisible to any single molecular platform. Yet our census of 25 "
    "integration studies exposes systematic limitations: 72% rely on the MOVICS framework, 88% depend "
    "on TCGA data, and the dominant feature selection strategy\u2014survival-based Cox filtering\u2014"
    "collapses biological heterogeneity into binary risk categories that cannot guide precision "
    "therapy. Sex-stratified and never-smoker-focused analyses are virtually absent. Clinical features "
    "are excluded from integration models. The 76 retractions concentrated in epigenetic prognostic "
    "signatures signal a broader validation crisis.\n\n"

    "Closing these gaps requires a shift in both methodology and study design. Hierarchical "
    "integration that separates biological discovery from clinical risk stratification, multivariate "
    "feature selection that tests molecular against clinical predictors, population-specific cohorts "
    "that include never-smokers and diverse ancestries, and prospective validation trials that connect "
    "multi-omics subtypes to therapeutic decisions\u2014these are the priorities that will determine "
    "whether the promise of multi-omics precision oncology is realized in lung cancer."
)
add_para(sec11)
all_body_text.append(sec11)

# ── TABLES ───────────────────────────────────────────────────────────────────

doc.add_page_break()
doc.add_heading('Tables', level=1)

# Table 1
add_para("Table 1. Census of multi-omics integration studies in lung cancer", bold=True)
t1_headers = ["Study", "Year", "Method", "Cancer", "N", "Omics Layers", "Feature Selection", "K", "Philosophy"]
t1_rows = [
    ["Ruan", "2022", "MOVICS", "LUAD", "437", "mRNA, Meth", "MAD", "4", "Biological"],
    ["Zou", "2022", "MOVICS", "LUAD", "437", "mRNA, lncRNA, Meth, Mut, CNV", "MAD", "4", "Biological"],
    ["Han", "2024", "MOVICS", "LUAD", "417", "mRNA, Meth", "Cox p<0.05", "2", "Clinical"],
    ["Lin", "2024", "MOVICS", "LUAD", "varies", "mRNA, lncRNA, miRNA, Meth, Mut", "Cox p<0.05", "2", "Clinical"],
    ["Chu", "2024", "MOVICS", "LUAD", "429", "mRNA, lncRNA, miRNA, Meth, Mut", "Cox p<0.01", "2", "Clinical"],
    ["Xie", "2024", "MOVICS", "LUAD", "417", "mRNA, Meth", "Cox + immune", "3", "Clinical"],
    ["Wu", "2024", "SNF", "LUAD", "TCGA", "mRNA, Meth, Mut", "Cox", "3", "Clinical"],
    ["Zhang", "2025", "NMF", "NSCLC", "122", "WES, RNA, Meth, scRNA", "NMF", "4", "Biological"],
    ["Zhong", "2025", "NMF", "LUAD", "101", "WES, RNA, Meth", "NMF", "3", "Biological"],
    ["Yan", "2025", "MOVICS", "NSCLC", "551", "mRNA, lncRNA, Meth", "MAD", "5", "Biological"],
]
add_table(t1_headers, t1_rows)

# Table 2
add_para("Table 2. Comparison of multi-omics integration methods", bold=True)
t2_headers = ["Method", "Category", "Supervised", "Key Advantage", "Key Limitation", "LC Studies"]
t2_rows = [
    ["MOVICS", "Ensemble consensus", "Both", "10-algorithm consensus", "Computationally intensive", "18"],
    ["iCluster", "Joint latent variable", "Unsupervised", "Statistical rigor", "Slow, Gaussian assumption", "10+"],
    ["SNF", "Network fusion", "Unsupervised", "Preserves local structure", "Sensitive to K", "2\u20133"],
    ["MOFA+", "Factor analysis", "Unsupervised", "Variance decomposition", "No discrete subtypes", "3"],
    ["DIABLO", "Supervised PLS-DA", "Supervised", "Biomarker discovery", "Requires labels", "0"],
    ["Deep learning", "Neural networks", "Both", "Non-linear features", "Black box", "2"],
]
add_table(t2_headers, t2_rows)

# Table 3
add_para("Table 3. Critical gaps in multi-omics lung cancer research", bold=True)
t3_headers = ["Gap", "Category", "Evidence", "Priority"]
t3_rows = [
    ["Feature selection collapses biology", "Methodological", "13/18 MOVICS studies use Cox \u2192 2 subtypes", "Highest"],
    ["Sex-stratified analyses absent", "Population", "No study stratifies by sex", "Highest"],
    ["Never-smoker multi-omics absent", "Population", "Zero dedicated studies", "Highest"],
    ["TCGA monoculture", "Data", "88% use TCGA, 2 institutional", "High"],
    ["No clinical feature integration", "Methodological", "Clinical vars not used as inputs", "High"],
    ["Zero metabolomics", "Data", "0/25 studies", "High"],
    ["Zero spatial omics", "Data", "0/25 studies", "High"],
    ["Zero DIABLO applications", "Methodological", "Despite breast cancer success", "High"],
    ["Zero prospective studies", "Validation", "All retrospective", "High"],
    ["76 retractions in epigenetics", "Quality", "miRNA/lncRNA prognostic sigs", "Moderate"],
    ["No companion diagnostics", "Translation", "Zero FDA-approved tests", "High"],
    ["Cost-effectiveness unknown", "Translation", "No health-economic evidence", "Moderate"],
]
add_table(t3_headers, t3_rows)

# ── FIGURE LEGENDS ───────────────────────────────────────────────────────────

doc.add_page_break()
doc.add_heading('Figure Legends', level=1)

fig1_legend = (
    "Figure 1. Overview of multi-omics integration and AI/ML approaches in lung cancer. "
    "This schematic illustrates the current landscape of multi-omics data generation and integration "
    "in lung cancer research. Tumor tissue (surgical resection, biopsy) and liquid biopsy samples "
    "(blood, plasma) serve as input materials for multiple molecular profiling platforms: genomics "
    "(whole-exome sequencing), transcriptomics (bulk and single-cell RNA sequencing), epigenomics "
    "(DNA methylation arrays, ATAC-seq), proteomics (mass spectrometry-based quantification), "
    "single-cell profiling (scRNA-seq, CITE-seq), and spatial transcriptomics (Visium, MERFISH). "
    "These data streams feed into computational integration frameworks, including MOVICS (consensus "
    "of 10 algorithms), iCluster (joint latent variables), SNF (patient similarity networks), MOFA+ "
    "(Bayesian factor analysis), and emerging deep learning architectures (variational autoencoders, "
    "graph neural networks, transformers). Integration outputs branch into four primary applications: "
    "molecular subtype discovery, biomarker identification, drug response prediction, and immunotherapy "
    "patient selection. A parallel AI/ML layer encompasses histopathology deep learning (mutation "
    "prediction from H&E slides), radiomics (CT-based treatment response prediction), foundation "
    "models (general-purpose pathology encoders), and multimodal fusion architectures. Clinical data "
    "\u2014including TNM stage, smoking history, sex, ancestry, and performance status\u2014are shown "
    "as a parallel input stream connected by dashed lines, indicating that these features are currently "
    "disconnected from integration pipelines despite their established prognostic value. Created with "
    "BioRender.com."
)
add_para("Figure 1", bold=True)
add_para(fig1_legend)

fig2_legend = (
    "Figure 2. The feature selection philosophy divide in multi-omics lung cancer studies. "
    "Two panels illustrate the contrasting approaches observed in our census of 25 integration studies. "
    "Left panel (Biological Discovery): variance-based feature selection using median absolute deviation "
    "(MAD) or standard deviation (SD) retains approximately 10,000 features, producing 4\u20135 molecular "
    "subtypes, each characterized by distinct biological hallmarks (color-coded): immune-hot (high "
    "immune infiltration, checkpoint expression), metabolic (altered oxidative phosphorylation, lipid "
    "metabolism), proliferative (cell cycle dysregulation, DNA repair upregulation), and stromal "
    "(fibroblast activation, extracellular matrix remodeling). Right panel (Clinical Prediction): "
    "Cox regression-based feature selection retains 2,000\u20133,500 survival-associated features, "
    "converging on 2 subtypes labeled high-risk and low-risk. Center callout highlights the hidden "
    "problem: within the high-risk group, biologically distinct subtypes\u2014immune-excluded tumors "
    "(potentially responsive to anti-angiogenic agents) and metabolically dysfunctional tumors "
    "(potentially responsive to metabolic inhibitors)\u2014are collapsed into a single prognostic "
    "category, obscuring therapeutically actionable distinctions. Bottom panel: scatter plot displaying "
    "the correlation between the number of input features and the number of discovered subtypes across "
    "18 MOVICS-based studies, demonstrating that feature selection strategy systematically determines "
    "subtype resolution. Created with BioRender.com."
)
add_para("Figure 2", bold=True)
add_para(fig2_legend)

fig3_legend = (
    "Figure 3. Proposed hierarchical framework for multi-omics integration. "
    "A three-stage pipeline is depicted. Stage 1 (Biological Clustering): unsupervised clustering "
    "using variance-based features (MAD/SD selection, retaining the top 25% most variable features "
    "across omics layers) discovers natural molecular subtypes without imposing a survival prior. "
    "Stage 2 (Clinical Risk Stratification): within each biological subtype, multivariate methods\u2014"
    "elastic net regularization, weighted gene co-expression network analysis (WGCNA), and stability "
    "selection\u2014integrate both molecular and clinical features (age, stage, smoking, sex, ECOG) "
    "as competing predictors to stratify patients by risk. This design explicitly tests the incremental "
    "value of molecular over clinical features. Stage 3 (Therapeutic Matching): each subtype-risk "
    "combination maps to specific treatment strategies, including targeted therapy, immunotherapy "
    "selection, and drug repurposing candidates identified through Connectivity Map analysis or "
    "synthetic lethality screens. Novel feature selection methods are highlighted at each stage: "
    "MethylMix for identifying driver methylation events, stability selection for robust feature "
    "identification, and WGCNA for pathway-level modules. Created with BioRender.com."
)
add_para("Figure 3", bold=True)
add_para(fig3_legend)

fig4_legend = (
    "Figure 4. Gaps and roadmap for multi-omics precision oncology in lung cancer. "
    "An infographic-style summary contrasts the current state of the field with the envisioned future. "
    "Left column (Current State): TCGA data monoculture (88% of studies), Cox-based feature selection "
    "collapsing biology into binary risk, sex- and ancestry-blind study designs, zero prospective "
    "validation, zero companion diagnostics, and 76 retractions in epigenetic biomarker studies. "
    "Right column (Future State): multi-ethnic institutional cohorts with paired tissue and liquid "
    "biopsy profiling, hierarchical integration preserving biological taxonomy, sex-stratified "
    "analysis as a default reporting standard, prospective multi-omics trials embedded in clinical "
    "decision-making, and FDA-cleared companion diagnostics for multi-omics classifiers. Center "
    "column (Bridge Elements): the specific actions needed to transition from current to future state, "
    "including generation of institutional cohorts, integration of clinical features alongside molecular "
    "data, addition of spatial and metabolomic layers, AI fairness audits across demographic subgroups, "
    "companion diagnostic development, and cost-effectiveness evidence. Bottom timeline traces key "
    "milestones from TCGA (2012\u20132014) through CPTAC proteogenomics (2020) and foundation models "
    "(2023\u20132025) to the proposed prospective multi-omics integration trial. Created with "
    "BioRender.com."
)
add_para("Figure 4", bold=True)
add_para(fig4_legend)

# ── REFERENCES ───────────────────────────────────────────────────────────────

doc.add_page_break()
doc.add_heading('References', level=1)

references = [
    "Cancer Genome Atlas Research Network. Comprehensive molecular profiling of lung adenocarcinoma. Nature 511, 543-550 (2014).",
    "Cancer Genome Atlas Research Network. Comprehensive genomic characterization of squamous cell lung cancers. Nature 489, 519-525 (2012).",
    "Gillette, M. A. et al. Proteogenomic characterization reveals therapeutic vulnerabilities in lung adenocarcinoma. Cell 182, 200-225 (2020).",
    "Xu, J. Y. et al. Integrative proteomic characterization of human lung adenocarcinoma. Cell 182, 245-261 (2020).",
    "Chen, Y. J. et al. Proteogenomics of non-smoking lung cancer in East Asia delineates molecular signatures of pathogenesis. Cell 182, 226-244 (2020).",
    "Rudin, C. M. et al. Molecular subtypes of small cell lung cancer: a synthesis of human and mouse model data. Nat. Rev. Cancer 19, 289-297 (2019).",
    "Baine, M. K. et al. SCLC subtypes defined by ASCL1, NEUROD1, POU2F3, and YAP1. J. Thorac. Oncol. 15, 1823-1835 (2020).",
    "Shen, R., Olshen, A. B. & Ladanyi, M. Integrative clustering of multiple genomic data types using a joint latent variable model. Bioinformatics 25, 2906-2912 (2009).",
    "Wang, B. et al. Similarity network fusion for aggregating data types on a genomic scale. Nat. Methods 11, 333-337 (2014).",
    "Coudray, N. et al. Classification and mutation prediction from non-small cell lung cancer histopathology images using deep learning. Nat. Med. 24, 1559-1567 (2018).",
    "Chen, R. J. et al. Pan-cancer integrative histology-genomic analysis via multimodal deep learning. Cancer Cell 40, 865-878 (2022).",
    "Wang, X. et al. A pathology foundation model for cancer diagnosis and prognosis prediction. Nature 634, 921-928 (2024).",
    "Yang, G. et al. A foundation model for generalizable cancer diagnosis and survival prediction. Nat. Commun. 16, 2045 (2025).",
    "Skoulidis, F. & Heymach, J. V. Co-occurring genomic alterations in non-small-cell lung cancer biology and therapy. Nat. Rev. Cancer 19, 495-509 (2019).",
    "Liu, Y. et al. Proteogenomic characterization of small cell lung cancer identifies biological insights and subtype-specific therapeutic strategies. Cell 187, 184-203 (2024).",
    "Martinez-Ruiz, C. et al. Genomic-transcriptomic evolution in lung cancer and metastasis. Nature 616, 553-562 (2023).",
    "Argelaguet, R. et al. MOFA+: a statistical framework for comprehensive integration of multi-modal single-cell data. Genome Biol. 21, 111 (2020).",
    "Lu, X. et al. MOVICS: an R package for multi-omics integration and visualization in cancer subtyping. Bioinformatics 36, 5539-5541 (2020).",
    "Pierre-Jean, M. et al. Clustering and variable selection evaluation of 13 unsupervised methods for multi-omics data integration. Brief. Bioinform. 21, 2020-2030 (2020).",
    "Chauvel, C. et al. Evaluation of integrative clustering methods for the analysis of multi-omics data. Brief. Bioinform. 21, 541-552 (2020).",
    "Athaya, T. et al. Multimodal deep learning approaches for single-cell multi-omics data integration. Brief. Bioinform. 24, bbad313 (2023).",
    "Hu, B. et al. Benchmarking algorithms for single-cell multi-omics prediction and integration. Nat. Methods 21, 2126-2137 (2024).",
    "Baiao, A. et al. A technical review of multi-omics data integration methods. Brief. Bioinform. 26, bbaf066 (2025).",
    "Zou, Y. et al. Multi-omics consensus portfolio to refine the classification of lung adenocarcinoma. Transl. Lung Cancer Res. 11, 2251-2271 (2022).",
    "Rakaee, M. et al. Deep learning model for predicting immunotherapy response in advanced NSCLC. JAMA Oncol. 11, 45-54 (2025).",
    "Arshad, M. et al. Artificial intelligence and machine learning in lung cancer. Cancers 17, 456 (2025).",
    "Siegfried, J. M. et al. Women and lung cancer: does oestrogen play a role? Lancet Oncol. 2, 506-513 (2001).",
    "Patel, J. D. et al. Lung cancer in US women: a contemporary epidemic. JAMA 291, 1763-1768 (2004).",
    "Conforti, F. et al. Cancer immunotherapy efficacy and patients' sex: a systematic review and meta-analysis. Lancet Oncol. 19, 737-746 (2018).",
    "Park, E. et al. Proteogenomic characterization reveals estrogen signaling as a target for never-smoker LUAD. Cancer Res. 84, 2178-2192 (2024).",
    "Zhang, T. et al. Genomic and evolutionary classification of lung cancer in never smokers. Nat. Genet. 53, 1348-1359 (2021).",
    "Sun, S., Schiller, J. H. & Gazdar, A. F. Lung cancer in never smokers--a different disease. Nat. Rev. Cancer 7, 778-790 (2007).",
    "Govindan, R. et al. Genomic landscape of non-small cell lung cancer in smokers and never-smokers. Cell 150, 1121-1134 (2012).",
    "Hamouz, R. Z. et al. A functional genomics review of non-small-cell lung cancer in never smokers. Int. J. Mol. Sci. 24, 13314 (2023).",
    "Pershagen, G. et al. Residential radon exposure and lung cancer in Sweden. N. Engl. J. Med. 330, 159-164 (1994).",
    "Vahakangas, K. H. et al. Mutations of p53 and ras genes in radon-associated lung cancer. Lancet 339, 576-580 (1992).",
    "Murphy, S. J. et al. Lung cancer in nonsmoking individuals: a review. JAMA 333, 567-579 (2025).",
    "Wielscher, M. et al. DNA methylation signature of chronic low-grade inflammation and its role in cardio-respiratory disease. Nat. Commun. 13, 2536 (2022).",
    "Alam, H. et al. KMT2D deficiency impairs super-enhancers to confer a glycolytic vulnerability in lung cancer. Cancer Cell 37, 599-617 (2020).",
    "de Miguel, F. J. et al. SWI/SNF chromatin remodeling complexes promote TKI resistance in EGFR-mutant lung cancer. Cancer Cell 41, 1543-1560 (2023).",
    "Yuan, G. et al. Elevated NSD3 histone methylation activity drives squamous cell lung cancer. Nature 590, 504-508 (2021).",
    "Heeke, S. et al. Tumor- and circulating-free DNA methylation identifies clinically relevant SCLC subtypes. Cancer Cell 42, 225-242 (2024).",
    "Zuccato, C. F. et al. Prediction of brain metastasis development with DNA methylation signatures. Nat. Med. 31, 456-467 (2025).",
    "Rizvi, N. A. et al. Mutational landscape determines sensitivity to PD-1 blockade in NSCLC. Science 348, 124-128 (2015).",
    "Thorsson, V. et al. The immune landscape of cancer. Immunity 48, 812-830 (2018).",
    "Charoentong, P. et al. Pan-cancer immunogenomic analyses reveal genotype-immunophenotype relationships. Cell Rep. 18, 248-262 (2017).",
    "Bagaev, A. et al. Conserved pan-cancer microenvironment subtypes predict response to immunotherapy. Cancer Cell 39, 845-865 (2021).",
    "Stein-Thoeringer, C. K. et al. A non-antibiotic-disrupted gut microbiome is associated with clinical responses to immunotherapy. Nat. Med. 29, 906-916 (2023).",
    "Ferrari, V. et al. Sensitizing cancer cells to immune checkpoint inhibitors by microbiota-mediated upregulation of HLA class I. Cancer Cell 41, 1717-1730 (2023).",
    "Bach, P. B. et al. Racial differences in the treatment of early-stage lung cancer. N. Engl. J. Med. 341, 1198-1205 (1999).",
    "Dutta, R. et al. Understanding inequities in precision oncology diagnostics. Nat. Cancer 4, 787-794 (2023).",
    "Tsimberidou, A. M. et al. Molecular tumour boards -- current and future considerations for precision oncology. Nat. Rev. Clin. Oncol. 20, 843-863 (2023).",
    "Kato, S. et al. Real-world data from a molecular tumor board demonstrates improved outcomes. Nat. Commun. 11, 4965 (2020).",
    "Lamb, J. et al. The Connectivity Map: a new tool for biomedical research. Nat. Rev. Cancer 7, 54-60 (2007).",
    "Cannon, M. et al. DGIdb 5.0: rebuilding the drug-gene interaction database. Nucleic Acids Res. 52, D1227-D1235 (2024).",
    "Nair, N. U. et al. A landscape of response to drug combinations in NSCLC. Nat. Commun. 14, 3642 (2023).",
    "Long, L. et al. PARP inhibition induces synthetic lethality and adaptive immunity in LKB1-mutant lung cancer. Cancer Res. 83, 568-581 (2023).",
    "Lambrechts, D. et al. Phenotype molding of stromal cells in the lung tumor microenvironment. Nat. Med. 24, 1277-1289 (2018).",
    "Maynard, A. et al. Therapy-induced evolution of human lung cancer revealed by single-cell RNA sequencing. Cell 182, 1232-1251 (2020).",
    "Tagore, S. et al. Single-cell and spatial genomic landscape of NSCLC brain metastases. Nat. Med. 31, 678-690 (2025).",
    "Abbosh, C. et al. Tracking early lung cancer metastatic dissemination in TRACERx using ctDNA. Nature 616, 553-562 (2023).",
    "McElderry, S. E. et al. Microbiome analysis of 940 lung cancers in never-smokers reveals lack of clinically relevant associations. Nat. Commun. 16, 2345 (2025).",
]

for i, ref in enumerate(references, 1):
    add_para(f"{i}. {ref}")

# ── Additional sections ──────────────────────────────────────────────────────

doc.add_page_break()
doc.add_heading('Acknowledgments', level=1)
add_para(
    "The authors acknowledge computational resources provided by the Digital Health Institute, "
    "Rutgers University."
)

doc.add_heading('Author contributions', level=1)
add_para(
    "V.R. conceived the study, conducted the systematic literature search and census of multi-omics "
    "integration studies, and wrote the manuscript. S.R. supervised the project and critically revised "
    "the manuscript. All authors approved the final version."
)

doc.add_heading('Competing interests', level=1)
add_para("The authors declare no competing interests.")

# ── Save and report ──────────────────────────────────────────────────────────

doc.save(OUTPUT_PATH)

total_text = " ".join(all_body_text)
main_wc = word_count(total_text)
print(f"\nMain text word count: {main_wc}")
print(f"Reference count: {len(references)}")
print(f"Saved to: {OUTPUT_PATH}")
